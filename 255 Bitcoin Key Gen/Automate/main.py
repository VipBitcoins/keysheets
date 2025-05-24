import subprocess
import re
import os
import shutil
import qrcode
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
import time
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from concurrent.futures import ThreadPoolExecutor, as_completed
import json

TEMP_FOLDER = "temp_qr_codes"
template_path = "template/template.docx"
def generate_single_vanity():
    # Path to vanitygen.exe (update this if needed)
    VANITYGEN_PATH = "./oclvanitygen/vanitygen.exe"
    VANITYGEN_COMMAND = "-i -v 1vip"

    # Regular expressions to extract Address and Privkey
    ADDRESS_REGEX = re.compile(r'Address:\s*(\S+)')
    PRIVKEY_REGEX = re.compile(r'Privkey:\s*(\S+)')

    """Runs vanitygen.exe once and extracts the address and private key."""
    try:
        process = subprocess.run(
            ["powershell", "-Command", f'& {VANITYGEN_PATH} {VANITYGEN_COMMAND}'],
            capture_output=True,
            text=True
        )
        
        output = process.stdout

        # Extract Address and Privkey
        address_match = ADDRESS_REGEX.search(output)
        privkey_match = PRIVKEY_REGEX.search(output)
        
        if address_match and privkey_match:
            return {
                "address": address_match.group(1),
                "privkey": privkey_match.group(1)
            }
    except Exception as e:
        print(f"Error running command: {e}")
    
    return None

def generate_vanity_addresses(num_addresses=255, num_threads=5):
    """Generates multiple vanity addresses using threading."""
    print(f"Generating {num_addresses} vanity addresses using {num_threads} threads...")

    results = []
    with ThreadPoolExecutor(max_workers=num_threads) as executor:
        # Submit tasks to the thread pool
        futures = [executor.submit(generate_single_vanity) for _ in range(num_addresses)]
        
        for future in as_completed(futures):
            result = future.result()
            if result:
                results.append(result)
            
            # Print progress
            if len(results) % 10 == 0:
                print(f"Generated {len(results)}/{num_addresses} addresses...")
    
    return results

from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def add_numbers_and_text(file_path, output_file, data_dict):
    """
    Adds numbers (1-255) and corresponding text from a dictionary to an existing table.
    Increases left & right padding inside table cells to allow better text wrapping.
    """
    doc = Document(file_path)

    # Font settings
    FONT_NAME = "Liberation Serif"
    FONT_SIZE = Pt(3)  # Font size for numbers
    FONT_SIZE_TEXT = Pt(2.9)  # Font size for text

    table = doc.tables[0]  # Assumes the first table in the document

    # Set cell margin (left and right padding)
    cell_margin_left = 200  # 200 twips (~0.18 inches)
    cell_margin_right = 200

    num = 1
    for row in table.rows:
        for cell in row.cells:
            if num <= 255:
                # Apply padding by modifying cell properties in XML
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcPr.append(parse_xml(
                    f'<w:tcMar {nsdecls("w")}><w:left w:w="{cell_margin_left}" w:type="dxa"/><w:right w:w="{cell_margin_right}" w:type="dxa"/></w:tcMar>'
                ))

                # Clear existing content
                paragraph = cell.paragraphs[0]
                paragraph.clear()
                
                # Align the number paragraph to the right
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                paragraph.paragraph_format.space_after = Pt(9)  # Padding before text

                # Add the number
                run_num = paragraph.add_run(str(num))
                run_num.font.name = FONT_NAME
                run_num.font.size = FONT_SIZE

                # Get text and enforce 5 lines of wrapping
                text_to_add = data_dict.get(num, "")

                # Add a new paragraph for text and align it centrally
                text_paragraph = cell.add_paragraph(text_to_add)
                text_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Apply bold formatting to text
                if text_to_add:
                    run_text = text_paragraph.runs[0]
                    # run_text.bold = True
                    run_text.font.name = FONT_NAME
                    run_text.font.size = FONT_SIZE_TEXT

                num += 1

    doc.save(output_file)
    print(f"Numbers and text added successfully! Saved as {output_file}")



def generate_qr_code(data, filename):
    """
    Generates a QR code image for the given data and saves it at `filename`,
    which is expected to be a full path.
    """
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=5,  # Adjust as needed
        border=2
    )
    qr.add_data(data)
    qr.make(fit=True)
    
    img = qr.make_image(fill="black", back_color="white")
    img.save(filename)  

def add_numbers_qr_codes(file_path, output_file, data_dict):
    """
    Adds numbers from 1 to 255, starting from the rightmost column of the first row
    and moving leftward. Places a QR code in each cell without changing the table structure.
    """

    # Create temp folder if it doesn't exist
    if not os.path.exists(TEMP_FOLDER):
        os.makedirs(TEMP_FOLDER)

    doc = Document(file_path)

    # Font settings
    FONT_NAME = "Liberation Serif"
    FONT_SIZE = Pt(3)  # Font size 3pt for both number and text

    # Assume the first table in the document
    table = doc.tables[0]

    # Counter for numbering (start from 1)
    num = 1

    for row in table.rows:
        for cell in reversed(row.cells):
            if num <= 255:
                # Number paragraph, right-aligned
                paragraph = cell.paragraphs[0]
                paragraph.clear()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                run_num = paragraph.add_run(str(num))
                run_num.font.name = FONT_NAME
                run_num.font.size = FONT_SIZE

                # Vertically center the cell
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # Generate the QR code file before inserting
                qr_filename = os.path.join(TEMP_FOLDER, f"qr_{num}.png")
                generate_qr_code(data_dict.get(num, ""), qr_filename)

                # Now add QR code in a new paragraph
                qr_paragraph = cell.add_paragraph()
                qr_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # Adjust vertical spacing if you want the QR code lower:
                qr_paragraph.paragraph_format.space_before = Pt(3)

                run_qr = qr_paragraph.add_run()
                run_qr.add_picture(qr_filename, width=Inches(0.19), height=Inches(0.19))

                num += 1

    # Save
    doc.save(output_file)
    print(f"Numbers and QR codes added successfully! Saved as {output_file}")

    # Cleanup
    shutil.rmtree(TEMP_FOLDER)
    print("Temporary QR codes deleted successfully.")

def save_public_addresses_to_excel(output_file, public_addresses):
    """Saves all public addresses into an Excel file."""
    print("Saving public addresses to Excel...")
    df = pd.DataFrame({"PUBLIC ADDRESSES": public_addresses})
    writer = pd.ExcelWriter(output_file, engine='xlsxwriter')
    df.to_excel(writer, index=False)
    worksheet = writer.sheets['Sheet1']
    worksheet.set_column('A:A', 50)  # Adjust column width
    writer.close()
    print(f"Public addresses saved to {output_file}")

def main():
    
    """Runs all steps sequentially."""
    print("Starting process...")
    results = generate_vanity_addresses()
    # Reload results from the JSON file
    # with open("results.json", "r") as json_file:
    #     results = json.load(json_file)
    # print("Results reloaded from results.json")
    private_keys = {i+1: res['privkey'] for i, res in enumerate(results)}
    public_addresses = [res['address'] for res in results]
    public_dict = {i+1: res['address'] for i, res in enumerate(results)}
    add_numbers_and_text(template_path,"private_keys.docx", private_keys)
    add_numbers_qr_codes(template_path,"public_qr_codes.docx", public_dict)
    save_public_addresses_to_excel("public_addresses.xlsx", public_addresses)
    print("Process completed successfully!")

if __name__ == "__main__":
    main()
